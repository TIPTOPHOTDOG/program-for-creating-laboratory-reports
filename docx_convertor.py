import csv
from docx import Document
from docx.shared import Pt
from docx.shared import Mm, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


class MyDocument:
    def print_data(self):
        data = {}
        with open('user_data.csv', 'r', newline='') as csvfile:
            csvreader = csv.reader(csvfile)
            headers = next(csvreader)
            for row in csvreader:
                for i in range(len(headers)):
                    data[headers[i]] = row[i]

        for key in data.keys():
            print(key, data[key])

    def create_lab_report(self):
        data = {}
        with open('user_data.csv', 'r', newline='') as csvfile:
            csvreader = csv.reader(csvfile)
            headers = next(csvreader)
            for row in csvreader:
                for i in range(len(headers)):
                    data[headers[i]] = row[i]
        print("Value of Field 1 from the second row:", data["Field 1"])

        doc = Document()
        if 'Normal' in doc.styles:
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(14)
        else:
            style = doc.styles.add_style('Normal')
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)

        style = doc.styles.add_style('BoldNormal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = True

        style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(12)

        doc.add_heading(f'{data["Field 1"]}', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph = doc.paragraphs[0]
        header_paragraph.style = 'Normal'

        section = doc.sections[0]
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.right_margin = Mm(20)
        for i in range(3):
            doc.add_paragraph("")

        doc.add_heading(f'{data["Field 2"]}', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_paragraph = doc.paragraphs[-1]
        header_paragraph.style = 'Normal'
        for i in range(5):
            doc.add_paragraph("")

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f'Отчет по лабораторной работе {data["Field 3"]}')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(24)
        paragraph.paragraph_format.space_after = Pt(1)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f'по дисциплине «{data["Field 4"]}»')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(16)

        for i in range(3):
            doc.add_paragraph("")

        # Информация о студенте
        paragraph = doc.add_paragraph('')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1 = paragraph.add_run(f'Выполнил: студент группы {data["Field 5"]}')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1.font.size = Pt(14)

        run1.add_break()

        run2 = paragraph.add_run(f' {data["Field 6"]}{"⠀" * 6}')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2.font.size = Pt(14)

        run2.add_break()

        run3 = paragraph.add_run(f'Проверил: {data["Field 7"]}{"⠀" * 6}')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3.font.size = Pt(14)

        for i in range(5):
            doc.add_paragraph("")

        paragraph = doc.add_paragraph(f'Тамбов {data["Field 8"]}')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = doc.add_heading(f'{"⠀" * 4}Цель работы:')
        paragraph.style = 'BoldNormal'

        paragraph.paragraph_format.space_after = Pt(1)

        doc.add_paragraph(f'{"⠀" * 4}{data["Field 11"]}')

        paragraph = doc.add_heading(f'{"⠀" * 4}Задание:', level=2)
        paragraph.style = 'BoldNormal'
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph = doc.add_paragraph(f'{"⠀" * 4}Вариант №{data["Field 9"]}')
        paragraph.paragraph_format.space_after = Pt(1)

        paragraph = doc.add_paragraph(f'{"⠀" * 4}{data["Field 12"]}')

        paragraph = doc.add_heading(f'{"⠀" * 4}Решение:', level=2)
        paragraph.style = 'BoldNormal'
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph = doc.add_paragraph(f'{"⠀" * 4}{data["Field 13"]}')
        paragraph.style = 'Normal'

        paragraph = doc.add_heading('Листинг программы', level=2)
        paragraph.style = 'BoldNormal'

        paragraph.paragraph_format.space_after = Pt(1)

        for i in data["Field 14"].split('\n'):
            code = doc.add_paragraph(i)
            code.style = "Code"
            code.paragraph_format.space_after = Pt(1)
            code.keep_together = True

        paragraph = doc.add_heading('Результаты работы программы', level=2)
        paragraph.style = 'BoldNormal'

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(data["Field 16"], width=Inches(6))
        run.add_text(f'Рисунок 2. {data["Field 17"]}')
        paragraph.style = 'Normal'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(data["Field 18"], width=Inches(6))
        run.add_text(f'Рисунок 2. {data["Field 19"]}')
        paragraph.style = 'Normal'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = doc.add_heading(f'{"⠀" * 4}Выводы:', level=2)
        paragraph.style = 'BoldNormal'
        paragraph = doc.add_paragraph(f'{"⠀" * 4}{data["Field 15"]}')
        paragraph.style = 'Normal'
        doc.save('lab_report.docx')
        return 1
